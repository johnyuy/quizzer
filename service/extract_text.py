import io
import pdfplumber
import fitz  # PyMuPDF
from docx import Document as DocxDocument

def extract_text(file_data, file_type: str) -> str:
    """
    Extract text from PDF, DOCX, or TXT files.

    Args:
        file_data: Uploaded file object or file path
        file_type: 'pdf', 'docx', 'txt'

    Returns:
        str: Extracted text
    """
    text = ""

    if file_type == "pdf":
        try:
            # First try PyMuPDF (fitz) for robust extraction
            pdf_doc = fitz.open(stream=file_data.read(), filetype="pdf")
            for page in pdf_doc:
                text += page.get_text()
            pdf_doc.close()
        except Exception as e1:
            # fallback to pdfplumber if fitz fails
            try:
                file_data.seek(0)  # reset pointer
                with pdfplumber.open(file_data) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text
            except Exception as e2:
                raise ValueError(f"Failed to read PDF. fitz error: {e1}, pdfplumber error: {e2}")

    elif file_type == "docx":
        try:
            if isinstance(file_data, io.BytesIO):
                doc = DocxDocument(file_data)
            else:
                doc = DocxDocument(file_data.name)
            for para in doc.paragraphs:
                text += para.text + "\n"
        except Exception as e:
            raise ValueError(f"Failed to read DOCX: {e}")

    elif file_type == "txt":
        try:
            if isinstance(file_data, io.BytesIO):
                file_data.seek(0)
                text = file_data.read().decode("utf-8")
            else:
                with open(file_data.name, "r", encoding="utf-8") as f:
                    text = f.read()
        except Exception as e:
            raise ValueError(f"Failed to read TXT: {e}")

    else:
        raise ValueError(f"Unsupported file type: {file_type}")

    return text.strip()